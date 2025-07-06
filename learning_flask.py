"""
from flask import Flask, request, send_file, jsonify
import pandas as pd
from docx import Document
import os

app = Flask(__name__)

def convert_excel_to_word(excel_path, word_path):
    try:
        df = pd.read_excel(excel_path)
        doc = Document()
        doc.add_heading('Excel Data', level=1)

        for _, row in df.iterrows():
            doc.add_paragraph(', '.join(map(str, row.values)))

        doc.save(word_path)
        return word_path
    except Exception as e:
        return str(e)

@app.route('/upload', methods=['GET'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files['file']
    
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400

    if not file.filename.endswith(('.xlsx', '.xls')):
        return jsonify({"error": "Invalid file type. Please upload an Excel file."}), 400

    try:
        excel_path = "temp.xlsx"
        file.save(excel_path)
        word_path = "converted.docx"
        result = convert_excel_to_word(excel_path, word_path)

        if os.path.exists(word_path):
            return send_file(word_path, as_attachment=True)
        else:
            return jsonify({"error": f"Conversion failed: {result}"}), 500

    except Exception as e:
        return jsonify({"error": f"Internal server error: {str(e)}"}), 500

if __name__ == '__main__':
    app.run(debug=True)


""

from flask import Flask, request, send_file
import os

app = Flask(__name__)

@app.route('/upload')
def upload_file():
    if 'file' not in request.files:
        return "No file uploaded!", 400

    file = request.files['file']
    if not file.filename.endswith(('.xlsx', '.xls')):
        return "Invalid file type!", 400

    excel_path = "/home/Ola/python_lab/FSI-2023-DOWNLOAD.xlsx"   #path to our excelfile to be converetd by this programe
    word_path = "./Downloads/converted.docx"
    file.save(excel_path)

    convert_excel_to_word(excel_path, word_path)
    return send_file(word_path, as_attachment=True)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)

""

from flask import Flask, render_template
#to input file into flask, we need two modules flask_wtf and wtforms to the installed using pip and import here for use as below
from flask_wtf import FlaskForm
from wtforms import FileField, SubmitField

app = Flask(__name__)
app.config['SECRET_KEY'] = 'supersecretkey'

#to build the form upload, we would crrate a class

class UploadFileForm(FlaskForm): #inherits from the imported Flaskform above
    file = FileField("File") #creating an input variables using the imported FileField above
    submit = SubmitField("Upload File") #creating an input variables using the imported submitField above.

@app.route('/', methods=["GET", "POST"])
@app.route('/home', methods=["GET", "POST"])
def home():
    form = UploadFileForm
    return render_template("index.html", form=form)


if __name__ == '__main__':
    app.run(debug=True)

"""