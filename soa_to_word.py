#pip install openpyxl python-docx


import pandas as pd

from docx import Document

soa = pd.read_excel("/content/drive/MyDrive/soa/CIVIS SOA.xlsx")


doc = Document()
doc.add_heading("Statement of Applicability", level=2)
doc.add_heading("The Risk Treatment Plan has been defined within the ISMS. The definition includes reference to the required headings within Annex A of the ISO 27001:2022 Standard. The Statement of Applicability was noted to have identified controls to mitigate risks following identification, analysis and evaluation and was evidenced. The following was observed:", level=6)

#print(f"The Risk Treatment Plan has been defined within the ISMS. The definition includes reference to the required headings within Annex A of the ISO 27001:2022 Standard. + '\n' + The Statement of Applicability was noted to have identified controls to mitigate risks following identification, analysis and evaluation and was evidenced. The following was observed:")

for index, row in soa.iterrows():
  doc.add_paragraph()
  for col_name, value in row.items():
    doc.add_paragraph(f"{col_name}: {value}")



doc.save("/content/drive/MyDrive/soa/SOAword2.docx")
print(f"Word document saved as {'/content/drive/MyDrive/soa/SOAword2.docx'}")